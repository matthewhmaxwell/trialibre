"""Word document (.docx) ingestor."""

from __future__ import annotations

import io
from pathlib import Path


class DocxIngestor:
    """Extract text from Word documents."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    @property
    def format_name(self) -> str:
        return "Word Document"

    async def extract_text(self, source: str | bytes, **kwargs) -> str:
        from docx import Document

        if isinstance(source, bytes):
            doc = Document(io.BytesIO(source))
        else:
            doc = Document(source)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    async def extract_structured(self, source: str | bytes, **kwargs) -> dict:
        return {}
